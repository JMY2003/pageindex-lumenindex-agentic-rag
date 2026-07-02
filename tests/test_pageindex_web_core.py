import asyncio
from pathlib import Path

from fastapi.testclient import TestClient
from pageindex_web.agent import AgentTools, react_event_stream
from pageindex_web.context_compact import compact_messages
from pageindex_web.indexer import _docx_pages_from_xml, _refine_large_pdf_nodes, index_document
from pageindex_web.main import INDEX_WORKERS, choose_nodes_with_llm, collect_pages, contextual_retrieval_question, conversation_context_from_messages, create_app, prepare_session_context, session_context_messages, usable_conversation_id, visible_session_messages
from pydantic import ValidationError
from pageindex_web.prompts import get_prompt, get_react_tools
from pageindex_web.search import SearchIndexCache, terms
from pageindex_web.settings import SettingsStore
from pageindex_web.storage import DocumentStore
from pageindex_web.main import password_hash, verify_password


def test_markdown_headings_build_nested_tree(tmp_path: Path):
    md = tmp_path / "guide.md"
    md.write_text(
        "# System Configuration\n\nOverview.\n\n## Cache Strategy\n\nSHA-256 fingerprints and gzip JSON.\n\n## Deep Reasoning\n\nInspect the outline before searching.\n",
        encoding="utf-8",
    )
    metadata, pages = index_document(md, lambda *_: None)

    assert metadata["index_strategy"] == "markdown_headings"
    assert metadata["page_count"] == len(pages)
    assert metadata["structure"][0]["title"] == "System Configuration"
    assert [node["title"] for node in metadata["structure"][0]["nodes"]] == ["Cache Strategy", "Deep Reasoning"]


def test_markdown_ignores_headings_inside_code_fences(tmp_path: Path):
    md = tmp_path / "code.md"
    md.write_text(
        "# Real Title\n\n```python\n# Not A Heading\nprint('ok')\n```\n\n## Real Section\n\nBody.\n",
        encoding="utf-8",
    )
    metadata, _ = index_document(md, lambda *_: None)

    titles = [metadata["structure"][0]["title"]] + [node["title"] for node in metadata["structure"][0]["nodes"]]
    assert titles == ["Real Title", "Real Section"]


def test_markdown_thins_fragmented_tiny_leaf_nodes(tmp_path: Path):
    md = tmp_path / "fragmented.md"
    tiny_sections = "\n".join(f"## Tiny {idx}\nx" for idx in range(30))
    md.write_text(f"# Root\n\nIntro.\n\n{tiny_sections}\n", encoding="utf-8")
    metadata, _ = index_document(md, lambda *_: None)

    assert metadata["index_strategy"] == "markdown_headings_thinned"
    assert metadata["structure"][0]["title"] == "Root"
    assert metadata["structure"][0]["nodes"] == []


def test_large_pdf_node_is_refined_from_layout_entries():
    pages = [{"page": idx, "content": f"Page {idx} content"} for idx in range(1, 21)]
    nodes = [{"node_id": "0001", "title": "Chapter", "start_index": 1, "end_index": 20, "pages": [1, 20], "summary": "", "nodes": []}]
    entries = [
        {"level": 1, "title": "Section A", "start": 2, "approximate": True},
        {"level": 1, "title": "Section B", "start": 9, "approximate": True},
        {"level": 1, "title": "Section C", "start": 15, "approximate": True},
    ]

    refined = _refine_large_pdf_nodes(nodes, pages, entries)

    assert refined == 1
    assert nodes[0]["refined_from_layout"] is True
    assert [child["title"] for child in nodes[0]["nodes"]] == ["Section A", "Section B", "Section C"]
    assert nodes[0]["nodes"][-1]["pages"] == [15, 20]


def test_search_cache_reuses_index_for_same_signature():
    docs = [
        {
            "id": "doc_a",
            "name": "a.md",
            "fingerprint": "sha256:a",
            "page_count": 1,
            "structure": [{"node_id": "0001", "title": "Cache Strategy", "pages": [1, 1], "summary": "SHA-256"}],
            "pages": [{"page": 1, "content": "The system uses SHA-256 fingerprints to identify duplicate documents."}],
        }
    ]
    cache = SearchIndexCache()
    first = cache.get(docs)
    second = cache.get(docs)

    assert first is second
    result = first.search("duplicate documents")[0]
    assert result["document_id"] == "doc_a"
    assert "proximity" in result
    assert "coverage" in result


def test_search_index_supports_multiple_documents():
    docs = [
        {
            "id": "doc_a",
            "name": "pricing.md",
            "fingerprint": "sha256:a",
            "page_count": 1,
            "structure": [{"node_id": "0001", "title": "Pricing", "pages": [1, 1], "summary": "enterprise plan pricing"}],
            "pages": [{"page": 1, "content": "The enterprise plan costs 6000 yuan per month."}],
        },
        {
            "id": "doc_b",
            "name": "security.md",
            "fingerprint": "sha256:b",
            "page_count": 1,
            "structure": [{"node_id": "0001", "title": "Security", "pages": [1, 1], "summary": "SOC2 audit controls"}],
            "pages": [{"page": 1, "content": "Security controls include SOC2 audit evidence and retention policies."}],
        },
    ]

    searcher = SearchIndexCache().get(docs)

    assert searcher.search("enterprise pricing", top_k=1)[0]["document_id"] == "doc_a"
    assert searcher.search("SOC2 audit", top_k=1)[0]["document_id"] == "doc_b"


def test_task_cancel_is_persisted(tmp_path: Path):
    store = DocumentStore(tmp_path)
    doc = store.new_document("a.md", 1, "text/markdown", "sha256:a", ".md")
    task = store.create_task(doc["id"])

    assert store.request_cancel(task.id)

    reloaded = DocumentStore(tmp_path)
    loaded = reloaded.get_task(task.id)
    assert loaded is not None
    assert loaded["cancelled"] is True
    assert loaded["status"] == "cancelled"


def test_sqlite_mirrors_documents_tasks_and_conversations(tmp_path: Path):
    store = DocumentStore(tmp_path)
    doc = store.new_document("a.md", 1, "text/markdown", "sha256:a", ".md")
    task = store.create_task(doc["id"])
    store.db.save_conversation(
        "conv_test",
        [doc["id"]],
        "standard",
        [{"role": "user", "content": "hello"}],
        "2026-01-01T00:00:00Z",
        "2026-01-01T00:00:00Z",
    )

    db_docs = store.db.list_documents()
    conversations = store.db.list_conversations()

    assert db_docs[0]["id"] == doc["id"]
    assert store.get_task(task.id)["id"] == task.id
    assert conversations[0]["id"] == "conv_test"
    assert conversations[0]["document_ids"] == [doc["id"]]

    assert store.db.delete_conversation("conv_test") is True
    assert store.db.get_conversation("conv_test") is None


def test_users_sessions_and_owner_scoped_documents(tmp_path: Path):
    store = DocumentStore(tmp_path)
    alice = store.db.create_user("alice", password_hash("correct horse"), "2026-01-01T00:00:00Z")
    bob = store.db.create_user("bob", password_hash("battery staple"), "2026-01-01T00:00:00Z")
    alice_doc = store.new_document("a.md", 1, "text/markdown", "sha256:a", ".md", alice["id"])
    bob_doc = store.new_document("b.md", 1, "text/markdown", "sha256:b", ".md", bob["id"])
    store.db.create_session(alice["id"], "token-a", "2026-01-01T00:00:00Z")

    assert verify_password("correct horse", store.db.get_user_by_username("alice")["password_hash"])
    assert store.db.get_session_user("token-a")["username"] == "alice"
    assert [doc["id"] for doc in store.documents(alice["id"])] == [alice_doc["id"]]
    assert store.get(bob_doc["id"], alice["id"]) is None


def test_auth_rejects_wrong_password_and_admin_assets_are_isolated(tmp_path: Path, monkeypatch):
    monkeypatch.setattr("pageindex_web.main.ROOT", tmp_path)
    app = create_app()
    admin_client = TestClient(app)
    bob_client = TestClient(app)

    normal_admin_name = bob_client.post("/api/auth/register", json={"username": "admin", "password": "admin123", "confirm_password": "admin123"})
    assert normal_admin_name.status_code == 400

    mismatch = bob_client.post("/api/auth/register", json={"username": "mismatch", "password": "firstpass", "confirm_password": "otherpass"})
    assert mismatch.status_code == 400
    assert mismatch.json()["detail"]["code"] == "password_mismatch"

    first_regular = bob_client.post("/api/auth/register", json={"username": "firstuser", "password": "firstpass", "confirm_password": "firstpass"})
    assert first_regular.status_code == 200
    assert first_regular.json()["user"]["is_admin"] is False
    bob_client.post("/api/auth/logout")

    admin_register = admin_client.post("/api/auth/register-admin", json={"username": "admin", "password": "admin123", "confirm_password": "admin123"})
    assert admin_register.status_code == 200
    assert admin_register.json()["user"]["is_admin"] is True
    admin_client.post("/api/auth/logout")

    wrong_password = admin_client.post("/api/auth/login", json={"username": "admin", "password": "admin124"})
    assert wrong_password.status_code == 401

    correct_password = admin_client.post("/api/auth/login", json={"username": "admin", "password": "admin123"})
    assert correct_password.status_code == 200
    assert correct_password.json()["user"]["is_admin"] is True

    bob_register = bob_client.post("/api/auth/register", json={"username": "bob", "password": "bob12345", "confirm_password": "bob12345"})
    assert bob_register.status_code == 200
    assert bob_register.json()["user"]["is_admin"] is False

    store = DocumentStore(tmp_path)
    admin_id = correct_password.json()["user"]["id"]
    bob_id = bob_register.json()["user"]["id"]
    admin_doc = store.new_document("admin.md", 1, "text/markdown", "sha256:admin", ".md", admin_id)
    bob_doc = store.new_document("bob.md", 1, "text/markdown", "sha256:bob", ".md", bob_id)
    store.db.save_conversation("conv_bob", [bob_doc["id"]], "standard", [{"role": "user", "content": "private"}], "2026-01-01T00:00:00Z", "2026-01-01T00:00:00Z", bob_id)

    assert bob_client.get("/api/admin/overview").status_code == 403
    bob_docs = bob_client.get("/api/documents").json()["documents"]
    assert [doc["id"] for doc in bob_docs] == [bob_doc["id"]]

    overview = admin_client.get("/api/admin/overview")
    assert overview.status_code == 200
    assert {doc["id"] for doc in overview.json()["documents"]} >= {admin_doc["id"], bob_doc["id"]}
    assert {conversation["id"] for conversation in overview.json()["conversations"]} >= {"conv_bob"}

    assert bob_client.delete(f"/api/admin/documents/{admin_doc['id']}").status_code == 403
    assert admin_client.delete(f"/api/admin/documents/{bob_doc['id']}").status_code == 200
    assert store.get(bob_doc["id"]) is None


def test_admin_creation_is_limited_to_three_accounts(tmp_path: Path, monkeypatch):
    monkeypatch.setattr("pageindex_web.main.ROOT", tmp_path)
    app = create_app()

    for idx in range(3):
        client = TestClient(app)
        response = client.post("/api/auth/register-admin", json={"username": f"admin{idx}", "password": "admin123", "confirm_password": "admin123"})
        assert response.status_code == 200
        assert response.json()["user"]["is_admin"] is True

    overflow = TestClient(app).post("/api/auth/register-admin", json={"username": "admin3", "password": "admin123", "confirm_password": "admin123"})
    assert overflow.status_code == 403
    assert overflow.json()["detail"]["code"] == "admin_limit_reached"


def test_admin_can_create_reset_and_delete_users_with_assets(tmp_path: Path, monkeypatch):
    monkeypatch.setattr("pageindex_web.main.ROOT", tmp_path)
    app = create_app()
    admin_client = TestClient(app)
    user_client = TestClient(app)

    admin_register = admin_client.post("/api/auth/register-admin", json={"username": "owneradmin", "password": "admin123", "confirm_password": "admin123"})
    assert admin_register.status_code == 200

    denied = user_client.post("/api/auth/register", json={"username": "regular", "password": "regular123", "confirm_password": "regular123"})
    assert denied.status_code == 200
    assert user_client.post("/api/admin/users", json={"username": "managed"}).status_code == 403

    created = admin_client.post("/api/admin/users", json={"username": "managed"})
    assert created.status_code == 200
    assert created.json()["user"]["is_admin"] is False
    assert created.json()["temporary_password"] == "managed@123"
    assert user_client.post("/api/auth/login", json={"username": "managed", "password": "managed@123"}).status_code == 200

    user_client.post("/api/auth/logout")
    custom_register = user_client.post("/api/auth/register", json={"username": "resetme", "password": "custom123", "confirm_password": "custom123"})
    assert custom_register.status_code == 200
    reset_user_id = custom_register.json()["user"]["id"]
    store = DocumentStore(tmp_path)
    store.db.save_conversation(
        "conv_resetme",
        [],
        "standard",
        [{"role": "user", "content": "keep my history"}],
        "2026-01-01T00:00:00Z",
        "2026-01-01T00:00:00Z",
        reset_user_id,
    )

    reset = admin_client.post(f"/api/admin/users/{reset_user_id}/reset-password")
    assert reset.status_code == 200
    assert reset.json()["temporary_password"] == "resetme@123"
    assert store.db.get_conversation("conv_resetme", reset_user_id) is not None
    assert user_client.post("/api/auth/login", json={"username": "resetme", "password": "custom123"}).status_code == 401
    assert user_client.post("/api/auth/login", json={"username": "resetme", "password": "resetme@123"}).status_code == 200

    managed_user = store.db.get_user_by_username("managed")
    managed_doc = store.new_document("managed.md", 1, "text/markdown", "sha256:managed", ".md", managed_user["id"])
    store.db.save_conversation(
        "conv_managed",
        [managed_doc["id"]],
        "standard",
        [{"role": "user", "content": "remove me"}],
        "2026-01-01T00:00:00Z",
        "2026-01-01T00:00:00Z",
        managed_user["id"],
    )

    deleted = admin_client.delete(f"/api/admin/users/{managed_user['id']}")
    assert deleted.status_code == 200
    assert deleted.json()["deleted_assets"] == {"documents": 1, "conversations": 1}
    assert store.get(managed_doc["id"]) is None
    assert store.db.get_conversation("conv_managed") is None
    assert store.db.get_user_by_username("managed") is None
    assert user_client.post("/api/auth/login", json={"username": "managed", "password": "managed@123"}).status_code == 401

    second_admin_client = TestClient(app)
    second_admin = second_admin_client.post("/api/auth/register-admin", json={"username": "secondadmin", "password": "admin123", "confirm_password": "admin123"})
    assert second_admin.status_code == 200
    second_admin_id = second_admin.json()["user"]["id"]
    second_admin_doc = store.new_document("second-admin.md", 1, "text/markdown", "sha256:second-admin", ".md", second_admin_id)
    store.db.save_conversation(
        "conv_second_admin",
        [second_admin_doc["id"]],
        "standard",
        [{"role": "user", "content": "admin private"}],
        "2026-01-01T00:00:00Z",
        "2026-01-01T00:00:00Z",
        second_admin_id,
    )

    reset_admin = admin_client.post(f"/api/admin/users/{second_admin_id}/reset-password")
    assert reset_admin.status_code == 403
    assert reset_admin.json()["detail"]["code"] == "admin_user_protected"

    delete_admin = admin_client.delete(f"/api/admin/users/{second_admin_id}")
    assert delete_admin.status_code == 403
    assert delete_admin.json()["detail"]["code"] == "admin_user_protected"

    delete_admin_doc = admin_client.delete(f"/api/admin/documents/{second_admin_doc['id']}")
    assert delete_admin_doc.status_code == 403
    assert delete_admin_doc.json()["detail"]["code"] == "admin_asset_protected"
    assert store.get(second_admin_doc["id"]) is not None

    delete_admin_conversation = admin_client.delete("/api/admin/conversations/conv_second_admin")
    assert delete_admin_conversation.status_code == 403
    assert delete_admin_conversation.json()["detail"]["code"] == "admin_asset_protected"
    assert store.db.get_conversation("conv_second_admin") is not None


def test_default_index_worker_target_matches_delivery_requirement():
    assert INDEX_WORKERS >= 5


def test_frontend_session_ids_are_accepted_by_backend():
    assert usable_conversation_id("conv_0123abcdEF_-") == "conv_0123abcdEF_-"
    assert usable_conversation_id("bad") is None
    assert usable_conversation_id("../conv_bad") is None


def test_context_compaction_has_deterministic_fallback(tmp_path: Path):
    settings = {"api_protocol": "openai", "api_key": ""}
    messages = []
    for idx in range(5):
        messages.extend([
            {"role": "user", "content": f"question {idx}"},
            {"role": "assistant", "content": f"answer {idx}"},
        ])

    compacted, info = compact_messages(settings, tmp_path / "logs", messages, keep_recent=2)

    assert info["compacted"] is True
    assert "Full pre-compaction context archive" in compacted[0]["content"]
    assert compacted[0]["role"] == "user"
    assert len(compacted) == 5


def test_conversation_context_preserves_session_messages_for_followups():
    messages = [
        {"role": "user", "content": "Summarize the pricing table."},
        {"role": "assistant", "content": "The second row covers the enterprise plan."},
        {"role": "user", "content": "What does that second row include?"},
    ]

    context = conversation_context_from_messages(messages)
    retrieval_question = contextual_retrieval_question("Compare it with the first row.", context)

    assert "User: Summarize the pricing table." in context
    assert "Assistant: The second row covers the enterprise plan." in context
    assert "Current question" in retrieval_question
    assert "Compare it with the first row." in retrieval_question
    assert "second row covers the enterprise plan" in retrieval_question


def test_session_context_uses_hidden_summary_and_recent_visible_messages():
    messages = [
        {"type": "context_summary", "content": "Earlier pricing discussion.", "covered_visible_count": 2},
        {"role": "user", "content": "old question"},
        {"role": "assistant", "content": "old answer"},
        {"role": "user", "content": "recent question"},
        {"role": "assistant", "content": "recent answer"},
    ]

    visible = visible_session_messages(messages)
    context = session_context_messages(messages)

    assert [item["content"] for item in visible] == ["old question", "old answer", "recent question", "recent answer"]
    assert context[0]["content"].startswith("Compressed session context")
    assert [item["content"] for item in context[1:]] == ["recent question", "recent answer"]


def test_prepare_session_context_compacts_persisted_session_without_hiding_visible_chat(tmp_path: Path):
    messages = []
    for idx in range(8):
        messages.extend([
            {"role": "user", "content": f"question {idx} " + ("x " * 260)},
            {"role": "assistant", "content": f"answer {idx} " + ("y " * 260)},
        ])

    session_messages, context, info = prepare_session_context(
        {"context_enabled": True, "context_window": 1000, "api_key": ""},
        tmp_path,
        messages,
    )

    assert info["compacted"] is True
    assert session_messages[0]["type"] == "context_summary"
    assert session_messages[0]["covered_visible_count"] > 0
    assert len(visible_session_messages(session_messages)) == len(messages)
    assert context[0]["content"].startswith("Compressed session context")
    assert "question 7" in conversation_context_from_messages(context)


def test_offline_react_stream_uses_real_tools():
    docs = [
        {
            "id": "doc_a",
            "name": "a.md",
            "fingerprint": "sha256:a",
            "page_count": 1,
            "structure": [{"node_id": "0001", "title": "Cache Strategy", "pages": [1, 1], "summary": "SHA-256"}],
            "pages": [{"page": 1, "content": "The system uses SHA-256 fingerprints to identify duplicate documents."}],
        }
    ]
    searcher = SearchIndexCache().get(docs)
    events = []

    async def collect():
        async for event in react_event_stream("duplicate documents", docs, searcher, {"step_budget": 10}, Path("/tmp"), lambda *_: []):
            events.append(event["event"])

    asyncio.run(collect())

    assert events.count("tool_call") == 3
    assert events[-1] == "final"


def test_offline_react_trace_includes_tool_reasons():
    docs = [
        {
            "id": "doc_a",
            "name": "a.md",
            "fingerprint": "sha256:a",
            "page_count": 1,
            "structure": [{"node_id": "0001", "title": "Cache Strategy", "pages": [1, 1], "summary": "SHA-256"}],
            "pages": [{"page": 1, "content": "The system uses SHA-256 fingerprints to identify duplicate documents."}],
        }
    ]
    searcher = SearchIndexCache().get(docs)
    tool_calls = []

    async def collect():
        async for event in react_event_stream("duplicate documents", docs, searcher, {"step_budget": 10}, Path("/tmp"), lambda *_: []):
            if event["event"] == "tool_call":
                tool_calls.append(event["data"])

    asyncio.run(collect())

    assert tool_calls[0]["reason"].startswith("Inspecting the outline")
    assert tool_calls[0]["tool"]["reason"] == tool_calls[0]["reason"]


def test_react_enforces_outline_before_search(monkeypatch):
    docs = [
        {
            "id": "doc_a",
            "name": "a.md",
            "fingerprint": "sha256:a",
            "page_count": 1,
            "structure": [{"node_id": "0001", "title": "Cache Strategy", "pages": [1, 1], "summary": "SHA-256"}],
            "pages": [{"page": 1, "content": "The system uses SHA-256 fingerprints."}],
        }
    ]
    searcher = SearchIndexCache().get(docs)

    class FakeToolCall:
        def __init__(self, call_id, name, arguments):
            self.id = call_id
            self.name = name
            self.arguments = arguments

    class FakeResponse:
        def __init__(self, text="", tool_calls=None):
            self.text = text
            self.tool_calls = tool_calls or []

    responses = iter([
        FakeResponse("I will search first.", [FakeToolCall("bad_search", "search", {"document_names": ["a.md"], "query": "fingerprints", "top_k": 3})]),
        FakeResponse("I will inspect the outline first.", [FakeToolCall("dir", "get_directory_structure", {"document_names": ["a.md"], "target": "ALL"})]),
        FakeResponse("Final answer from collected evidence."),
    ])

    monkeypatch.setattr("pageindex_web.agent.has_api_key", lambda _: True)
    monkeypatch.setattr("pageindex_web.agent.completion", lambda *_, **__: next(responses))
    events = []

    async def collect():
        async for event in react_event_stream("fingerprints", docs, searcher, {"api_key": "x", "step_budget": 10}, Path("/tmp"), lambda *_: []):
            events.append(event)

    asyncio.run(collect())

    assert any(event["event"] == "status" and event["data"].get("title") == "Inspecting outline first" for event in events)
    executed_tools = [event["data"]["tool"]["name"] for event in events if event["event"] == "tool_call"]
    assert executed_tools == ["get_directory_structure"]
    assert events[-1]["event"] == "final"


def test_public_tool_logic_exposes_directory_search_and_content():
    docs = [
        {
            "id": "doc_a",
            "name": "a.md",
            "fingerprint": "sha256:a",
            "page_count": 2,
            "structure": [
                {
                    "node_id": "0001",
                    "title": "Cache Strategy",
                    "pages": [1, 2],
                    "summary": "SHA-256 fingerprints",
                    "nodes": [{"node_id": "0002", "title": "Compression", "pages": [2, 2], "summary": "gzip pages"}],
                }
            ],
            "pages": [
                {"page": 1, "content": "The system uses SHA-256 fingerprints."},
                {"page": 2, "content": "Page bodies are cached as gzip JSON."},
            ],
        }
    ]
    searcher = SearchIndexCache().get(docs)
    tools = AgentTools(docs, searcher, collect_pages)

    directory = tools.get_directory_structure(["a.md"], "ALL")
    results = tools.search(["a.md"], "gzip cache", 4)
    pages = tools.get_page_content(["a.md"], [2])

    assert directory[0]["structure"][0]["node_id"] == "0001"
    assert results[0]["document_id"] == "doc_a"
    assert pages[0]["page"] == 2
    assert "gzip JSON" in pages[0]["content"]


def test_agent_tool_arguments_are_pydantic_validated():
    docs = [
        {
            "id": "doc_a",
            "name": "a.md",
            "fingerprint": "sha256:a",
            "page_count": 1,
            "structure": [{"node_id": "0001", "title": "Cache Strategy", "pages": [1, 1], "summary": "SHA-256"}],
            "pages": [{"page": 1, "content": "The system uses SHA-256 fingerprints."}],
        }
    ]
    tools = AgentTools(docs, SearchIndexCache().get(docs), collect_pages)

    assert tools.call("search", {"document_names": ["a.md"], "query": "SHA-256", "top_k": 1})[0]["document_id"] == "doc_a"
    try:
        tools.call("search", {"document_names": ["a.md"], "query": "SHA-256", "top_k": 999})
    except ValidationError:
        pass
    else:
        raise AssertionError("invalid tool arguments should fail Pydantic validation")


def test_llm_node_selection_is_pydantic_validated(monkeypatch):
    docs = [
        {
            "id": "doc_a",
            "name": "a.md",
            "fingerprint": "sha256:a",
            "page_count": 1,
            "structure": [{"node_id": "0001", "title": "Cache Strategy", "pages": [1, 1], "summary": "SHA-256"}],
            "pages": [{"page": 1, "content": "The system uses SHA-256 fingerprints."}],
        }
    ]

    class FakeResponse:
        text = """[
          {"document_id": "doc_a", "node_id": "0001", "reason": "matches"},
          {"document_id": "doc_a", "node_id": "", "reason": "empty node id"},
          {"document_id": "other", "node_id": "0001", "reason": "wrong document"}
        ]"""

    monkeypatch.setattr("pageindex_web.main.has_api_key", lambda _: True)
    monkeypatch.setattr("pageindex_web.main.completion", lambda *_, **__: FakeResponse())

    selected = choose_nodes_with_llm({"api_key": "x", "max_output_tokens": 1200}, "fingerprints", docs)

    assert selected == [{"document_id": "doc_a", "node_id": "0001", "reason": "matches"}]


def test_cache_hit_requires_current_split_cache(tmp_path: Path):
    store = DocumentStore(tmp_path)
    doc = store.new_document("a.md", 1, "text/markdown", "sha256:a", ".md")

    assert store.find_by_fingerprint("sha256:a") is None

    store.save_index(
        doc["id"],
        {"structure": [], "summary": "", "version": "fastapi-web-2", "cache_schema_version": 2, "page_count": 1},
        [{"page": 1, "content": "hello"}],
    )

    assert store.find_by_fingerprint("sha256:a")["id"] == doc["id"]


def test_settings_support_delivery_fields(tmp_path: Path):
    saved = SettingsStore(tmp_path / "settings.json").save({
        "context_window_k": 200,
        "max_output_tokens": 40000,
        "deep_thinking": False,
        "context_enabled": True,
    })

    assert saved["context_window"] == 200000
    assert saved["context_window_k"] == 200
    assert saved["max_output_tokens"] == 40000
    assert saved["deep_thinking"] is False
    assert saved["context_enabled"] is True


def test_app_settings_override_system_environment(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("OPENAI_API_KEY", "from-system")
    monkeypatch.setenv("OPENAI_API_BASE", "https://system.example/v1")
    store = SettingsStore(tmp_path / "settings.json")
    store.save({"api_protocol": "openai", "api_key": "from-app", "api_url": "https://app.example/v1", "model": "app-model"})

    loaded = store.load(include_secret=True)

    assert loaded["api_key"] == "from-app"
    assert loaded["api_url"] == "https://app.example/v1"
    assert loaded["model"] == "app-model"


def test_system_environment_is_fallback_without_app_settings(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("OPENAI_API_KEY", "from-system")
    monkeypatch.setenv("OPENAI_API_BASE", "https://system.example/v1")
    monkeypatch.setenv("PAGEINDEX_MODEL", "system-model")

    loaded = SettingsStore(tmp_path / "settings.json").load(include_secret=True)

    assert loaded["api_key"] == "from-system"
    assert loaded["api_url"] == "https://system.example/v1"
    assert loaded["model"] == "system-model"


def test_qwen_zshrc_style_environment_is_openai_compatible_fallback(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("QWEN_API_KEY", "qwen-key")
    monkeypatch.setenv("QWEN_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1")
    monkeypatch.setenv("QWEN_MODEL", "qwen-max")

    loaded = SettingsStore(tmp_path / "settings.json").load(include_secret=True)

    assert loaded["api_protocol"] == "openai"
    assert loaded["api_key"] == "qwen-key"
    assert loaded["api_url"] == "https://dashscope.aliyuncs.com/compatible-mode/v1"
    assert loaded["model"] == "qwen-max"


def test_zshrc_file_is_openai_compatible_fallback(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("HOME", str(tmp_path))
    for key in [
        "OPENAI_API_KEY",
        "CHATGPT_API_KEY",
        "QWEN_API_KEY",
        "DASHSCOPE_API_KEY",
        "OPENAI_API_BASE",
        "OPENAI_BASE_URL",
        "QWEN_BASE_URL",
        "QWEN_API_BASE",
        "DASHSCOPE_BASE_URL",
        "DASHSCOPE_API_BASE",
        "PAGEINDEX_MODEL",
        "QWEN_MODEL",
        "DASHSCOPE_MODEL",
    ]:
        monkeypatch.delenv(key, raising=False)
    (tmp_path / ".zshrc").write_text(
        "export QWEN_API_KEY='from-zshrc'\n"
        "export QWEN_BASE_URL=https://dashscope.aliyuncs.com/compatible-mode/v1\n"
        "export QWEN_MODEL=qwen-zshrc\n",
        encoding="utf-8",
    )

    loaded = SettingsStore(tmp_path / "settings.json").load(include_secret=True)

    assert loaded["api_protocol"] == "openai"
    assert loaded["api_key"] == "from-zshrc"
    assert loaded["api_url"] == "https://dashscope.aliyuncs.com/compatible-mode/v1"
    assert loaded["model"] == "qwen-zshrc"


def test_model_prompts_are_loaded_from_yaml():
    assert "LumenIndex" in get_prompt("react.system")
    assert "Question: cache strategy" in get_prompt("standard_answer.user", question="cache strategy", evidence="Evidence A")
    assert [tool["name"] for tool in get_react_tools()] == ["get_directory_structure", "search", "get_page_content"]


def test_cjk_question_terms_are_cleaned_and_expanded():
    extracted = terms("\u914d\u7f6e\u7ba1\u7406\u662f\u4ec0\u4e48\u610f\u601d\uff0c\u5305\u542b\u54ea\u4e9b\u7b56\u7565\uff1f")

    assert "\u914d\u7f6e\u7ba1\u7406" in extracted
    assert "\u4ec0\u4e48" not in extracted
    assert "\u914d\u7f6e" in extracted
    assert "\u7ba1\u7406" in extracted


def test_docx_xml_page_break_extraction(tmp_path: Path):
    from docx import Document

    path = tmp_path / "pages.docx"
    doc = Document()
    doc.add_paragraph("\u7b2c\u4e00\u9875")
    doc.add_page_break()
    doc.add_paragraph("\u7b2c\u4e8c\u9875")
    doc.save(path)

    pages = _docx_pages_from_xml(path)

    assert pages is not None
    assert len(pages) >= 2
    assert "\u7b2c\u4e00\u9875" in pages[0]["content"]
    assert "\u7b2c\u4e8c\u9875" in pages[1]["content"]
